import json
import os
import re
import time
import pandas as pd
from datetime import datetime
from playwright.sync_api import sync_playwright
from docx import Document
from docx.shared import Pt, RGBColor

# === 全局配置 ===
URL_HOME = "https://petermoportfolio.com/"
URL_MEMOS = "https://petermoportfolio.com/memos"

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DATA_DIR = os.path.join(BASE_DIR, "data")
HOLDINGS_DIR = os.path.join(DATA_DIR, "holdings")
MEMOS_DIR = os.path.join(DATA_DIR, "memos")

# 超时设置 (100秒)
TIMEOUT_MS = 100000 

os.makedirs(HOLDINGS_DIR, exist_ok=True)
os.makedirs(MEMOS_DIR, exist_ok=True)

def format_date_filename(text):
    """提取并格式化日期"""
    match = re.search(r'(\d{4})年(\d{1,2})月(\d{1,2})日', text)
    if match:
        year, month, day = match.groups()
        return f"{year}年{int(month):02d}月{int(day):02d}日"
    return None

def save_memo_doc(title, date_str, content):
    """保存 Word 文档"""
    try:
        filename_date = format_date_filename(title)
        if not filename_date:
            filename_date = format_date_filename(date_str)
        
        final_filename = f"{filename_date}.docx" if filename_date else f"{title}.docx"
        final_filename = re.sub(r'[\\/*?:"<>|]', "_", final_filename)

        doc = Document()
        doc.add_heading(title, 0)
        
        p_date = doc.add_paragraph(date_str)
        if p_date.runs:
            run_date = p_date.runs[0]
            run_date.italic = True
            run_date.font.color.rgb = RGBColor(128, 128, 128)
        
        doc.add_paragraph("") 
        if content:
            doc.add_paragraph(content)
        
        file_path = os.path.join(MEMOS_DIR, final_filename)
        doc.save(file_path)
        print(f"    [保存成功] {final_filename}")
        
    except Exception as e:
        print(f"    [保存失败] {title}: {e}")

def parse_holdings():
    """抓取持仓"""
    print("\n>>> [1/2] 正在启动抓取持仓...")
    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True) 
        page = browser.new_page()
        
        try:
            page.goto(URL_HOME, timeout=TIMEOUT_MS)
            page.wait_for_selector("div.text-muted-foreground", timeout=TIMEOUT_MS)
        except Exception as e:
            print(f"持仓页加载失败: {e}")
            browser.close()
            return

        holdings = []
        ticker_elements = page.locator("div.text-xs.text-muted-foreground").all()
        
        for ticker_el in ticker_elements:
            ticker_text = ticker_el.inner_text().strip()
            if len(ticker_text) > 8 or not ticker_text: continue
            try:
                name = ticker_el.locator("xpath=..//span[contains(@class, 'font-semibold')]").inner_text()
            except: name = "Unknown"
            share = "N/A"
            try:
                row_text = ticker_el.locator("xpath=../..").inner_text()
                match = re.search(r'(\d+\.?\d*%)', row_text)
                if match: share = match.group(1)
            except: pass

            holdings.append({
                "代码": ticker_text,
                "名称": name,
                "仓位": share,
                "抓取日期": datetime.now().strftime("%Y-%m-%d")
            })

        browser.close()
        
        if holdings:
            df = pd.DataFrame(holdings).drop_duplicates(subset=['代码'])
            today = datetime.now().strftime("%Y-%m-%d")
            df.to_excel(os.path.join(HOLDINGS_DIR, f"holdings_{today}.xlsx"), index=False)
            df.to_excel(os.path.join(HOLDINGS_DIR, "latest_holdings.xlsx"), index=False)
            print(f"    持仓数据已保存到 {HOLDINGS_DIR}")

def parse_memos_detailed():
    """
    【修复版】Memos 抓取逻辑
    策略：每次抓取前都强制刷新列表页，确保 DOM 状态最新，防止点击错位。
    """
    print("\n>>> [2/2] 正在启动笔记深度抓取 (稳定性增强模式)...")
    
    with sync_playwright() as p:
        # headless=False 方便观察，稳定后可改为 True
        browser = p.chromium.launch(headless=False) 
        context = browser.new_context()
        page = context.new_page()
        
        try:
            # 1. 首次访问获取总数
            page.goto(URL_MEMOS, timeout=TIMEOUT_MS)
            print("    等待列表加载...")
            page.wait_for_selector("div[data-slot='card']", timeout=TIMEOUT_MS)
            
            # 获取总篇数
            total_count = page.locator("div[data-slot='card']").count()
            print(f"    共发现 {total_count} 篇笔记，准备逐一抓取...")
            
            # 2. 循环处理
            for i in range(total_count):
                print(f"\n    正在处理第 {i+1}/{total_count} 篇...")
                
                try:
                    # === 核心修复 ===
                    # 每次循环都重新加载页面，保证索引(i)对应的就是第i个元素
                    # 不要依赖 go_back()，直接 goto
                    if i > 0: # 第一篇已经在页面上了，不用刷
                        print("    正在刷新列表，确保定位准确...")
                        page.goto(URL_MEMOS, timeout=TIMEOUT_MS)
                        page.wait_for_selector("div[data-slot='card']", timeout=TIMEOUT_MS)
                        # 稍微等一小会儿让布局稳定
                        page.wait_for_timeout(2000)

                    # 重新获取卡片列表 (因为页面刷新了，旧的变量失效了)
                    cards = page.locator("div[data-slot='card']")
                    current_card = cards.nth(i)
                    
                    # 滚动并点击
                    current_card.scroll_into_view_if_needed(timeout=10000)
                    current_card.click()
                    
                    # 等待详情页
                    page.wait_for_selector("div.prose", timeout=15000)
                    
                    # 提取数据
                    title = page.locator("div[data-slot='card-title']").first.inner_text().strip()
                    
                    date_str = ""
                    try:
                        date_str = page.locator("div[data-slot='card-title'] + p").first.inner_text().strip()
                    except: pass
                        
                    content = page.locator("div.prose").first.inner_text().strip()
                    
                    save_memo_doc(title, date_str, content)
                    
                except Exception as e:
                    print(f"    [错误] 处理第 {i+1} 篇时出错: {e}")
                    # 如果这篇出错了，继续循环下一篇
                    continue
                    
        except Exception as e:
            print(f"    [致命错误] 脚本中断: {e}")
            
        browser.close()

if __name__ == "__main__":
    parse_holdings()
    parse_memos_detailed()
    print("\n" + "="*50)
    print(f"全部任务完成！")
    print(f"持仓表格: {HOLDINGS_DIR}")
    print(f"笔记文档: {MEMOS_DIR}")
    print("="*50)